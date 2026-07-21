import pathlib
from docx import Document

# Absolute paths
md_path = pathlib.Path(r'C:/Users/DELL8/.gemini/antigravity/brain/21a6608e-8164-43af-a8d4-37a70b418fe7/roadmap_supplier_risk.md')
output_path = pathlib.Path(r'C:/Users/DELL8/OneDrive/Desktop/General/Cyber/roadmap_supplier_risk.docx')

# Read markdown content
md_text = md_path.read_text(encoding='utf-8')

# Simple conversion: headings start with #, tables start with |, otherwise paragraphs

doc = Document()

lines = md_text.splitlines()
for line in lines:
    stripped = line.strip()
    if not stripped:
        doc.add_paragraph('')
        continue
    if stripped.startswith('#'):
        level = len(stripped) - len(stripped.lstrip('#'))
        heading = stripped.lstrip('#').strip()
        doc.add_heading(heading, level=level)
    elif stripped.startswith('|'):
        # Add table row as a paragraph (simple representation)
        doc.add_paragraph(stripped)
    else:
        doc.add_paragraph(stripped)

doc.save(output_path)
print('Generated', output_path)
