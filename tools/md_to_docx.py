import pathlib, re
from docx import Document

# Paths
md_path = pathlib.Path(r'roadmap_supplier_risk.md')
output_path = pathlib.Path(r'roadmap_supplier_risk.docx')

# Read markdown
md_text = md_path.read_text(encoding='utf-8')

# Simple conversion: treat headings and tables
doc = Document()

lines = md_text.splitlines()
for line in lines:
    # Heading
    if line.startswith('#'):
        level = len(line) - len(line.lstrip('#'))
        heading = line.lstrip('#').strip()
        if level == 1:
            doc.add_heading(heading, level=0)
        else:
            doc.add_heading(heading, level=level)
    elif line.startswith('|'):
        # Simple handling: add as a table row (assume previous line was header)
        # We'll accumulate rows until a blank line
        # Use a flag to create table when first encountered
        # For simplicity, skip complex markdown tables and just add as preformatted text
        doc.add_paragraph(line)
    elif line.strip() == '':
        doc.add_paragraph('')
    else:
        doc.add_paragraph(line)

doc.save(output_path)
print('Generated', output_path)
